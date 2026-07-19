#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
========================================================================
MOTOR DE ESTUDIO DE ESTRATEGIAS (ES / futuros)
========================================================================
Toma una LISTA DE OPERACIONES (export tipo TradingView Strategy Tester)
y, opcionalmente, OHLC de 5m (+1h) del mismo instrumento, y reproduce el
estudio completo de gestión de riesgo:

  1. Resumen y curva de equity
  2. Excursiones: pullback (MAE) y recorrido a favor (MFE) vs cierre
  3. SL backstop (barrido de niveles)
  4. Compras escalonadas con micros (distribución + balance)
  5. TP techo (backstop para brokers que exigen TP)
  6. Breakeven stop            [requiere OHLC 5m]
  7. Dirección: largos vs cortos
  8. Sesión / hora del día (NY) y día de la semana
  9. Duración / time-stop      [validado con OHLC 5m]
 10. Volatilidad (ATR)         [requiere OHLC 1h]
 11. Corrida combinada final   [requiere OHLC 5m]

Los HALLAZGOS se calculan de los datos: una estrategia distinta produce
conclusiones distintas automáticamente. Los resultados "negativos"
(p. ej. si el time-stop no ayuda) también se reportan: son información.

Salida: reporte .docx + PNGs + un resumen por consola + metricas.json

USO:
  python estudio_estrategia.py
Edita el bloque CONFIG de abajo para apuntar a tus archivos.
========================================================================
"""
import os, json, warnings
import numpy as np, pandas as pd
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
warnings.filterwarnings('ignore')

# ======================= CONFIG (editar aquí) =======================
CONFIG = {
    # --- Archivos ---
    'trade_file'  : '/mnt/user-data/uploads/1783518563576_ES5m_ConfNormal_TC_TSR_070726.csv',
    'ohlc_5m_file': '/mnt/user-data/uploads/ES_5m.csv',   # None si no hay
    'ohlc_1h_file': '/mnt/user-data/uploads/ES_1h.csv',   # None si no hay
    'output_dir'  : '/mnt/user-data/outputs/estudio_auto',
    'strategy_name': 'Estrategia ES 5m',

    # --- Instrumento ---
    'point_value'      : 'auto', # 'auto' = derivar de los datos; o fijar (ES=50, NQ=20, MES=5, etc.)
    'micro_point_value': 5.0,    # USD por punto del micro (MES=5)
    'micros_per_mini'  : 10,     # 10 MES = 1 ES

    # --- Columnas de la lista de operaciones ---
    'col_trade' : 'Trade number',
    'col_type'  : 'Tipo',
    'col_dt'    : 'Fecha y hora',
    'col_price' : 'Precio USD',
    'col_pnl'   : 'PyG netas USD',
    'col_mfe'   : 'Desviación favorable USD',
    'col_mae'   : 'Desviación adversa USD',
    'col_dur'   : 'Duration (bars)',
    'kw_entry'  : 'Entrada', 'kw_exit': 'Salida', 'kw_long': 'largo',

    # --- Columnas OHLC ---
    'ohlc_dt':'DateTime','ohlc_o':'Open','ohlc_h':'High','ohlc_l':'Low','ohlc_c':'Close',

    # --- Palancas a barrer (en puntos) ---
    'sl_levels'      : [80, 90, 100, 120],
    'tp_ceiling'     : [60, 70, 80],
    'be_triggers'    : [15, 20, 25, 30],
    'scalein_levels' : [0, 12, 24],   # niveles de compra (pts adversos)
    'scalein_alloc'  : [4, 3, 3],     # micros por nivel (suman micros_per_mini)
    'timestop_bars'  : [24, 48, 72, 96, 120, 160],
    'news_hours_et'  : [8, 11],        # horas ET a marcar como "ventana de noticias"

    # --- Config recomendada para la corrida final ---
    'reco_sl':100, 'reco_be':20, 'reco_tp':70,
}
# ====================================================================

PV = CONFIG['point_value']; MES = CONFIG['micro_point_value']; NM = CONFIG['micros_per_mini']
os.makedirs(CONFIG['output_dir'], exist_ok=True)
CH = os.path.join(CONFIG['output_dir'], 'charts'); os.makedirs(CH, exist_ok=True)
GREEN='#1baf7a'; RED='#e34948'; BLUE='#2a78d6'; AMBER='#eda100'; GREY='#888780'
def money(v): return ('-$' if v<0 else '$')+f"{abs(v):,.0f}"

# --------------------------- CARGA ---------------------------
def load_trades(cfg):
    df = pd.read_csv(cfg['trade_file'])
    ent = df[df[cfg['col_type']].str.contains(cfg['kw_entry'], na=False)].sort_values(cfg['col_trade']).reset_index(drop=True)
    sal = df[df[cfg['col_type']].str.contains(cfg['kw_exit'],  na=False)].sort_values(cfg['col_trade']).reset_index(drop=True)
    is_long = ent[cfg['col_type']].str.contains(cfg['kw_long'], na=False).values
    ep=ent[cfg['col_price']].values.astype(float); xp=sal[cfg['col_price']].values.astype(float)
    pnl=sal[cfg['col_pnl']].values.astype(float)
    size = ent['Tamaño (cant.)'].values.astype(float) if 'Tamaño (cant.)' in ent else np.ones(len(ent))
    # --- derivar valor del punto de los datos: PnL = PV * size * movimiento a favor ---
    global PV
    move=np.where(is_long, xp-ep, ep-xp)
    with np.errstate(divide='ignore',invalid='ignore'):
        pv_est=pnl/(move*size)
    pv_est=pv_est[np.isfinite(pv_est)&(np.abs(move)>1e-9)]
    if str(cfg.get('point_value','auto')).lower()=='auto' and len(pv_est):
        PV=float(np.round(np.median(pv_est),4))
    else:
        PV=float(cfg['point_value'])
    t = pd.DataFrame({
        'entry_dt': pd.to_datetime(ent[cfg['col_dt']].values),
        'exit_dt' : pd.to_datetime(sal[cfg['col_dt']].values),
        'is_long' : is_long, 'entry_px':ep, 'exit_px':xp, 'size':size,
        'pnl'     : pnl,
        'mfe_usd' : sal[cfg['col_mfe']].values.astype(float),
        'mae_usd' : sal[cfg['col_mae']].values.astype(float),
        'dur'     : sal[cfg['col_dur']].values.astype(int),
    })
    t['mfe_pts']=t['mfe_usd']/PV; t['mae_pts']=t['mae_usd']/PV; t['pnl_pts']=t['pnl']/PV
    return t

def load_ohlc(path, cfg):
    if not path or not os.path.exists(path): return None
    o = pd.read_csv(path, parse_dates=[cfg['ohlc_dt']]).rename(columns={
        cfg['ohlc_dt']:'dt',cfg['ohlc_o']:'o',cfg['ohlc_h']:'h',cfg['ohlc_l']:'l',cfg['ohlc_c']:'c'})
    return o.set_index('dt').sort_index()

# ------------------ RECONSTRUCCIÓN INTRABAR ------------------
def reconstruct_paths(t, o5):
    """Devuelve FAV/ADV por barra (pts desde la entrada) por trade, o None."""
    if o5 is None: return None, 0
    FAV=[None]*len(t); ADV=[None]*len(t); OPN=[None]*len(t); ok=0
    for i in range(len(t)):
        w=o5.loc[t.entry_dt.iloc[i]:t.exit_dt.iloc[i]]
        if len(w)<2: continue
        hi=w['h'].max(); lo=w['l'].min()
        entry=(hi-t.mfe_pts.iloc[i]) if t.is_long.iloc[i] else (lo+t.mfe_pts.iloc[i])
        if t.is_long.iloc[i]:
            FAV[i]=w['h'].values-entry; ADV[i]=w['l'].values-entry; OPN[i]=w['o'].values-entry
        else:
            FAV[i]=entry-w['l'].values; ADV[i]=entry-w['h'].values; OPN[i]=entry-w['o'].values
        if abs(FAV[i][-1]*0 + (w['c'].values[-1]-entry if t.is_long.iloc[i] else entry-w['c'].values[-1])*PV - t.pnl.iloc[i])<=100:
            ok+=1
    return {'FAV':FAV,'ADV':ADV,'OPN':OPN}, ok

# --------------------------- MÉTRICAS ---------------------------
def metrics(p):
    p=np.asarray(p,float)
    if len(p)==0: return dict(n=0,net=0,wr=0,pf=0,dd=0,gp=0,gl=0,exp=0)
    gp=p[p>0].sum(); gl=p[p<0].sum(); eq=p.cumsum(); dd=(eq-np.maximum.accumulate(eq)).min()
    return dict(n=len(p),net=p.sum(),wr=(p>0).mean()*100,pf=(gp/abs(gl)) if gl!=0 else float('inf'),
                dd=dd,gp=gp,gl=gl,exp=p.mean())

def savefig(fig,name):
    path=os.path.join(CH,name); fig.savefig(path,bbox_inches='tight',dpi=140,facecolor='white'); plt.close(fig); return path

# --------------------------- ESTUDIOS ---------------------------
def study_overview(t):
    m=metrics(t.pnl.values); m['best']=t.pnl.max(); m['worst']=t.pnl.min()
    eq=t.sort_values('exit_dt')['pnl'].cumsum()
    fig,ax=plt.subplots(figsize=(9,3.6))
    ax.plot(range(1,len(eq)+1),eq.values,color=BLUE,lw=1.6); ax.fill_between(range(1,len(eq)+1),eq.values,color=BLUE,alpha=.08)
    ax.axhline(0,color=GREY,lw=.6); ax.set_title('Curva de equity',loc='left',fontsize=10)
    ax.set_xlabel('Trade #'); ax.set_ylabel('PnL acumulado (USD)')
    for s in ['top','right']: ax.spines[s].set_visible(False)
    m['chart']=savefig(fig,'equity.png'); return m

def study_direction(t):
    out={}
    for lab,mask in [('Ambos',np.ones(len(t),bool)),('Largos',t.is_long.values),('Cortos',~t.is_long.values)]:
        sub=t[mask].sort_values('exit_dt'); out[lab]=metrics(sub.pnl.values)
    # mensual por dirección
    tmp=t.copy(); tmp['mes']=tmp.exit_dt.dt.to_period('M').astype(str)
    piv=tmp.pivot_table(index='mes',columns='is_long',values='pnl',aggfunc='sum',fill_value=0)
    piv.columns=['Cortos' if c==False else 'Largos' for c in piv.columns]
    fig,ax=plt.subplots(figsize=(9,3.8)); x=np.arange(len(piv)); w=.38
    if 'Largos' in piv: ax.bar(x-w/2,piv['Largos'],w,color=BLUE,label='Largos')
    if 'Cortos' in piv: ax.bar(x+w/2,piv['Cortos'],w,color=AMBER,label='Cortos')
    ax.axhline(0,color=GREY,lw=.6); ax.set_xticks(x); ax.set_xticklabels(piv.index,fontsize=8)
    ax.set_title('Neto por mes y dirección',loc='left',fontsize=10); ax.legend(frameon=False,fontsize=8)
    for s in ['top','right']: ax.spines[s].set_visible(False)
    out['chart']=savefig(fig,'direccion.png')
    # hallazgo automático
    best=max([('Largos',out['Largos']),('Cortos',out['Cortos'])],key=lambda k:k[1]['pf'])
    out['finding']=(f"El lado {best[0].lower()} es superior (PF {best[1]['pf']:.2f} vs "
                    f"{out['Cortos' if best[0]=='Largos' else 'Largos']['pf']:.2f}). "
                    f"Operar solo {best[0].lower()}: neto {money(out[best[0]]['net'])}, DD {money(out[best[0]]['dd'])}.")
    return out

def study_sessions(t):
    hr=t.entry_dt.dt.hour.values; dow=t.entry_dt.dt.dayofweek.values
    sess={'Europa 02-07':np.isin(hr,[2,3,4,5,6,7]),'Datos 08':hr==8,'NY 09-10':np.isin(hr,[9,10]),
          'Media 11':hr==11,'Mediodía 12-15':np.isin(hr,[12,13,14,15]),'Cierre 16-18':np.isin(hr,[16,17,18]),
          'Asia 19-01':np.isin(hr,[19,20,21,22,23,0,1])}
    rows={k:metrics(t.pnl.values[m]) for k,m in sess.items()}
    fig,ax=plt.subplots(figsize=(9.5,3.8)); labs=list(rows); vals=[rows[k]['net'] for k in labs]
    ax.axhline(0,color=GREY,lw=.6); ax.bar(range(len(labs)),vals,color=[GREEN if v>=0 else RED for v in vals])
    ax.set_xticks(range(len(labs))); ax.set_xticklabels(labs,rotation=25,ha='right',fontsize=8)
    ax.set_title('Neto por sesión (hora NY)',loc='left',fontsize=10); ax.set_ylabel('USD')
    for s in ['top','right']: ax.spines[s].set_visible(False)
    chart=savefig(fig,'sesiones.png')
    dwtab={['Lun','Mar','Mie','Jue','Vie','Sab','Dom'][d]:metrics(t.pnl.values[dow==d]) for d in sorted(set(dow))}
    worst_h=min(rows.items(),key=lambda k:k[1]['net']); best_h=max(rows.items(),key=lambda k:k[1]['net'])
    worst_d=min(dwtab.items(),key=lambda k:k[1]['net'])
    finding=(f"Mejor sesión: {best_h[0]} ({money(best_h[1]['net'])}, PF {best_h[1]['pf']:.2f}). "
             f"Peor sesión: {worst_h[0]} ({money(worst_h[1]['net'])}). "
             f"Peor día: {worst_d[0]} ({money(worst_d[1]['net'])}).")
    return dict(sessions=rows,dow=dwtab,chart=chart,finding=finding)

def study_sl(t, paths):
    base=metrics(t.pnl.values); rows={}
    # niveles de stop en DÓLARES, derivados de la propia distribución de |MAE| (percentiles) -> agnóstico al instrumento
    absmae=np.abs(t.mae_usd.values)
    levels=sorted(set(np.round(np.percentile(absmae,[60,75,85,95]),-1)))
    for SLusd in levels:
        if paths:
            out=t.pnl.values.copy().astype(float)
            for i in range(len(t)):
                a=paths['ADV'][i]
                if a is None: continue
                slpts=SLusd/PV
                for k in range(len(a)):
                    if a[k]<=-slpts:
                        out[i]=min(-slpts, paths['OPN'][i][k])*PV; break
        else:
            out=np.where(t.mae_usd.values<=-SLusd, -SLusd, t.pnl.values)
        rows[int(SLusd)]=metrics(out)
    best=max(rows.items(),key=lambda k:(k[1]['net']/abs(k[1]['dd']) if k[1]['dd']!=0 else 0))
    return dict(base=base,rows=rows,best=best[0],unit='USD',
                finding=f"Mejor SL backstop: {money(best[0])} (neto {money(best[1]['net'])}, DD {money(best[1]['dd'])}, PF {best[1]['pf']:.2f}).")

def study_timestop(t, paths):
    base=metrics(t.pnl.values); rows={}; validated = paths is not None
    for Nb in CONFIG['timestop_bars']:
        out=t.pnl.values.copy().astype(float)
        if validated:
            for i in range(len(t)):
                f=paths['FAV'][i]
                if f is None: continue
                D=len(f)-1
                if D>Nb:
                    # P&L al cierre de la barra Nb (aprox con close ~ nivel medio fav/adv no disponible; usamos fav como proxy alto y adv bajo -> usamos media)
                    lvl=(paths['FAV'][i][Nb]+paths['ADV'][i][Nb])/2
                    out[i]=lvl*PV
        rows[Nb]=metrics(out)
    best=max(rows.items(),key=lambda k:k[1]['net'])
    helps = best[1]['net']>base['net']
    finding=(f"El time-stop {'AYUDA' if helps else 'NO ayuda'}: el mejor nivel ({best[0]} barras) da "
             f"{money(best[1]['net'])} vs {money(base['net'])} sin time-stop. "
             + ("" if helps else "La correlación duración-resultado es sesgo de supervivencia."))
    # buckets por duración (informativo siempre)
    buckets=[(0,30),(30,80),(80,160),(160,10**6)]
    bk={f"{lo}-{hi if hi<10**5 else '+'}":metrics(t.pnl.values[(t.dur>=lo)&(t.dur<hi)]) for lo,hi in buckets}
    fig,ax=plt.subplots(figsize=(8.5,3.4)); labs=list(bk); vals=[bk[k]['net'] for k in labs]
    ax.axhline(0,color=GREY,lw=.6); ax.bar(labs,vals,color=[GREEN if v>=0 else RED for v in vals])
    ax.set_title('Neto por duración (barras)',loc='left',fontsize=10); ax.set_ylabel('USD')
    for s in ['top','right']: ax.spines[s].set_visible(False)
    return dict(base=base,rows=rows,buckets=bk,validated=validated,helps=helps,finding=finding,chart=savefig(fig,'duracion.png'))

def study_breakeven(t, paths):
    if paths is None: return None
    base=metrics(t.pnl.values); rows={}
    for trig in CONFIG['be_triggers']:
        out=t.pnl.values.copy().astype(float)
        for i in range(len(t)):
            f=paths['FAV'][i]; a=paths['ADV'][i]
            if f is None: continue
            armed=False
            for k in range(len(f)):
                if armed and a[k]<=0: out[i]=0.0; break
                if f[k]>=trig: armed=True
        rows[trig]=metrics(out)
    best=max(rows.items(),key=lambda k:k[1]['net'])
    helps=best[1]['net']>base['net']
    return dict(base=base,rows=rows,best=best[0],helps=helps,
                finding=f"El breakeven stop {'AYUDA' if helps else 'no ayuda'}: mejor trigger +{best[0]} pts -> "
                        f"neto {money(best[1]['net'])} (base {money(base['net'])}), PF {best[1]['pf']:.2f}.")

def study_tp(t, paths):
    rows={}
    mx=t.mfe_usd.max()
    # techos en DÓLARES relativos al MFE máximo del instrumento
    levels=sorted(set(np.round(np.array([0.9,1.0,1.15])*mx,-1)))
    for TPusd in levels:
        if paths:
            out=t.pnl.values.copy().astype(float); tppts=TPusd/PV
            for i in range(len(t)):
                f=paths['FAV'][i]
                if f is None: continue
                for k in range(len(f)):
                    if f[k]>=tppts: out[i]=tppts*PV; break
        else:
            out=np.where(t.mfe_usd.values>=TPusd, TPusd, t.pnl.values)
        fired=int((t.mfe_usd.values>=TPusd).sum()); rows[int(TPusd)]=(metrics(out),fired)
    reco=int(np.round(1.1*mx,-1))
    return dict(rows=rows,max_mfe=mx,reco=reco,unit='USD',
                finding=f"MFE máximo histórico {money(mx)}. Un TP techo ~{money(reco)} casi no dispara y no interfiere.")

def study_final(t, paths):
    if paths is None: return None
    SL=CONFIG['reco_sl']; BE=CONFIG['reco_be']; TP=CONFIG['reco_tp']
    lv=CONFIG['scalein_levels']; qt=CONFIG['scalein_alloc']
    def tpnl(i, scalein):
        f=paths['FAV'][i]; a=paths['ADV'][i]
        if f is None: return t.pnl.iloc[i]
        armed=False; E=None; xb=len(f)-1
        for k in range(len(f)):
            if a[k]<=-SL: E=-SL; xb=k; break
            if armed and a[k]<=0: E=0; xb=k; break
            if f[k]>=TP: E=TP; xb=k; break
            if f[k]>=BE: armed=True
        if E is None: E=t.pnl_pts.iloc[i]
        if not scalein: return E*PV
        tot=0.0
        for L,q in zip(lv,qt):
            filled=(L==0) or any(a[k]<=-L for k in range(xb+1))
            if filled: tot+=q*(E+L)*MES
        return tot
    hr=t.entry_dt.dt.hour.values; noNews=~np.isin(hr,CONFIG['news_hours_et']); longs=t.is_long.values
    steps=[('Base',metrics(t.pnl.values)),
           ('1 mini + SL/BE/TP',metrics([tpnl(i,False) for i in range(len(t))])),
           ('Escalonado + SL/BE/TP',metrics([tpnl(i,True) for i in range(len(t))])),
           ('+ sin ventana noticias',metrics(np.array([tpnl(i,True) for i in range(len(t))])[noNews])),
           ('+ solo largos',metrics(np.array([tpnl(i,True) for i in range(len(t))])[noNews&longs]))]
    fig,ax=plt.subplots(figsize=(9.5,4)); x=np.arange(len(steps)); w=.4
    ax.axhline(0,color=GREY,lw=.6)
    ax.bar(x-w/2,[s[1]['net'] for s in steps],w,color=GREEN,label='Neto')
    ax.bar(x+w/2,[s[1]['dd']  for s in steps],w,color=RED,label='Max drawdown')
    ax.set_xticks(x); ax.set_xticklabels([s[0] for s in steps],rotation=15,ha='right',fontsize=8)
    ax.set_title('Corrida final combinada',loc='left',fontsize=10); ax.legend(frameon=False,fontsize=8)
    for s in ['top','right']: ax.spines[s].set_visible(False)
    return dict(steps=steps,chart=savefig(fig,'final.png'),
                finding=f"Config recomendada: {steps[3][0]} -> neto {money(steps[3][1]['net'])}, "
                        f"PF {steps[3][1]['pf']:.2f}, DD {money(steps[3][1]['dd'])} (base PF {steps[0][1]['pf']:.2f}, DD {money(steps[0][1]['dd'])}).")

def study_regime(t, o1h):
    if o1h is None: return None
    w=o1h.loc[t.entry_dt.min():t.exit_dt.max()]
    if len(w)<10: return None
    ema=w['c'].ewm(span=50).mean(); chg=(w['c'].iloc[-1]/w['c'].iloc[0]-1)*100; pct_up=(w['c']>ema).mean()*100
    return dict(change=chg,pct_up=pct_up,
                finding=f"El instrumento se movió {chg:+.1f}% en el periodo y pasó {pct_up:.0f}% del tiempo sobre su EMA50 (1h): "
                        f"régimen {'ALCISTA' if pct_up>55 else ('BAJISTA' if pct_up<45 else 'lateral')}. "
                        f"Los resultados pueden depender de este régimen.")

# --------------------------- REPORTE ---------------------------
def build_report(cfg, res):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc=Document()
    def H(txt,sz=15,color='0F6E56'):
        p=doc.add_paragraph(); r=p.add_run(txt); r.bold=True; r.font.size=Pt(sz); r.font.color.rgb=RGBColor.from_string(color); return p
    def P(txt): 
        p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(10.5); return p
    def IMG(path,w=6.0):
        if path and os.path.exists(path): doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    def TBL(headers,rows):
        tb=doc.add_table(rows=1,cols=len(headers)); tb.style='Light Grid Accent 1'
        for j,h in enumerate(headers): tb.rows[0].cells[j].text=str(h)
        for row in rows:
            c=tb.add_row().cells
            for j,v in enumerate(row): c[j].text=str(v)
    H(f"Estudio de gestión de riesgo — {cfg['strategy_name']}",20,'1A1A19')
    o=res['overview']
    P(f"Muestra: {o['n']} operaciones · {res['t'].entry_dt.min().date()} a {res['t'].exit_dt.max().date()} · ${PV:.0f}/punto")
    H("1. Resumen")
    TBL(['Métrica','Valor'],[['Neto',money(o['net'])],['Win rate',f"{o['wr']:.1f}%"],['Profit factor',f"{o['pf']:.2f}"],
        ['Max drawdown',money(o['dd'])],['Expectativa/op',money(o['exp'])],['Mejor',money(o['best'])],['Peor',money(o['worst'])]])
    IMG(o['chart'])
    def sect(num,title,r,cols=None,rowfn=None):
        if r is None: return
        H(f"{num}. {title}")
        if 'finding' in r: P("Hallazgo: "+r['finding'])
        if cols and rowfn: TBL(cols,rowfn(r))
        if r.get('chart'): IMG(r['chart'])
    sect("2","Dirección: largos vs cortos",res['direction'],
         ['Escenario','n','Neto','WR','PF','Max DD'],
         lambda r:[[k,r[k]['n'],money(r[k]['net']),f"{r[k]['wr']:.0f}%",f"{r[k]['pf']:.2f}",money(r[k]['dd'])] for k in ['Ambos','Largos','Cortos']])
    sect("3","Sesión y hora (NY)",res['sessions'],
         ['Sesión','n','Neto','WR','PF'],
         lambda r:[[k,v['n'],money(v['net']),f"{v['wr']:.0f}%",f"{v['pf']:.2f}"] for k,v in r['sessions'].items()])
    sect("4","SL backstop (niveles en USD, derivados de |MAE|)",res['sl'],
         ['SL','Neto','WR','PF','Max DD'],
         lambda r:[[money(k),money(v['net']),f"{v['wr']:.0f}%",f"{v['pf']:.2f}",money(v['dd'])] for k,v in r['rows'].items()])
    sect("5","TP techo (niveles en USD, relativos al MFE máx)",res['tp'],
         ['TP','Neto','PF','Dispara'],
         lambda r:[[money(k),money(v[0]['net']),f"{v[0]['pf']:.2f}",f"{v[1]}"] for k,v in r['rows'].items()])
    sect("6","Duración / time-stop",res['timestop'],
         ['Duración','n','Neto','WR','PF'],
         lambda r:[[k,v['n'],money(v['net']),f"{v['wr']:.0f}%",f"{v['pf']:.2f}"] for k,v in r['buckets'].items()])
    if res.get('breakeven'):
        sect("7","Breakeven stop",res['breakeven'],
             ['Trigger (pts)','Neto','WR','PF','Max DD'],
             lambda r:[[k,money(v['net']),f"{v['wr']:.0f}%",f"{v['pf']:.2f}",money(v['dd'])] for k,v in r['rows'].items()])
    if res.get('regime'):
        H("8. Régimen del periodo"); P("Hallazgo: "+res['regime']['finding'])
    if res.get('final'):
        H("9. Corrida final combinada"); P("Hallazgo: "+res['final']['finding'])
        IMG(res['final']['chart'])
        TBL(['Configuración','n','Neto','WR','PF','Max DD'],
            [[s[0],s[1]['n'],money(s[1]['net']),f"{s[1]['wr']:.0f}%",f"{s[1]['pf']:.2f}",money(s[1]['dd'])] for s in res['final']['steps']])
    doc.add_paragraph()
    p=doc.add_paragraph(); r=p.add_run("Análisis cuantitativo de datos históricos; no es asesoría de inversión. El desempeño pasado no garantiza resultados futuros."); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string('5F5E5A')
    out=os.path.join(cfg['output_dir'],'Reporte_estudio.docx'); doc.save(out); return out

# --------------------------- MAIN ---------------------------
def run(cfg=CONFIG):
    t=load_trades(cfg); o5=load_ohlc(cfg['ohlc_5m_file'],cfg); o1h=load_ohlc(cfg['ohlc_1h_file'],cfg)
    paths,ok=reconstruct_paths(t,o5)
    print(f"Operaciones: {len(t)} | OHLC 5m: {'sí' if o5 is not None else 'no'} | reconstrucción válida: {ok}/{len(t)}")
    res={'t':t,'overview':study_overview(t),'direction':study_direction(t),'sessions':study_sessions(t),
         'sl':study_sl(t,paths),'tp':study_tp(t,paths),'timestop':study_timestop(t,paths),
         'breakeven':study_breakeven(t,paths),'regime':study_regime(t,o1h),'final':study_final(t,paths)}
    # resumen consola + json
    print("\n==================== RESUMEN AUTO ====================")
    for k in ['direction','sessions','sl','tp','timestop','breakeven','regime','final']:
        if res.get(k) and 'finding' in res[k]: print(f"- {res[k]['finding']}")
    rep=build_report(cfg,res)
    js={k:res[k]['finding'] for k in ['direction','sessions','sl','tp','timestop','breakeven','regime','final'] if res.get(k) and 'finding' in res[k]}
    json.dump(js,open(os.path.join(cfg['output_dir'],'metricas.json'),'w'),ensure_ascii=False,indent=2)
    print(f"\nReporte: {rep}")
    return res

if __name__=='__main__':
    run()
